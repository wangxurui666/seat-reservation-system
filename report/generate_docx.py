# -*- coding: utf-8 -*-
"""将课程设计报告.md转换为Word文档"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re, os

doc = Document()

# ============================================================
# 样式设置
# ============================================================
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# 页面设置 A4
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============================================================
# Helper Functions
# ============================================================
def add_heading_styled(text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(13)
    return h

def add_para(text, bold=False, indent=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else None
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_bullet(text, level=0):
    """添加列表项"""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    return p

def add_code(text):
    """添加代码块（灰色背景等宽字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # 背景色
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_table(headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 深蓝色背景
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A3C34" w:val="clear"/>')
        tcPr.append(shd)

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10.5)

    doc.add_paragraph()  # 表后空行
    return table

def page_break():
    doc.add_page_break()

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("数据库系统概论")
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x34)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("课程设计报告")
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(38)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x34)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("自习室座位预约系统")
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(30)
run.font.color.rgb = RGBColor(0xC4, 0x93, 0x3C)

for _ in range(4):
    doc.add_paragraph()

for line in ["学院：_____________", "专业：_____________", "学号：_____________", "姓名：_____________"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(16)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2026年6月")
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(16)

# ============================================================
# 一、系统定义
# ============================================================
page_break()
add_heading_styled("一、系统定义", 1)

add_heading_styled("1. 系统定义", 2)
add_para("本系统是一个基于 Web 的自习室/图书馆座位预约管理平台，旨在解决高校自习座位资源紧张、占座现象普遍的问题。系统为在校学生提供在线浏览、预约、签到自习座位的完整服务，并为管理员提供自习室与座位的管理功能。", indent=True)
add_para("核心价值：通过引入“预约—签到—信用”闭环机制和数据库层面的并发控制，确保座位资源的公平分配和高效利用。", indent=True)

add_heading_styled("2. 系统说明", 2)
add_para("目标用户：", bold=True)
add_bullet("学生用户：注册后可浏览座位地图，按日期和时段预约空闲座位，并按时签到使用")
add_bullet("管理员用户：可管理自习室和座位信息，查看违约记录和统计数据")

add_para("核心业务流程：", bold=True)
for i, step in enumerate([
    "注册与登录：用户填写用户名、密码、真实姓名和学号完成注册，登录后进入系统",
    "浏览座位：按楼栋→自习室→座位三个层次查看，选择日期和时段后实时显示各座位可用状态",
    "预约座位：点击空闲座位的“预约”按钮，系统通过数据库行级锁（SELECT ... FOR UPDATE）确保同一座位同一时段仅一人预约成功",
    "签到确认：预约成功后需在 15 分钟内在线签到，超时未签到则自动标记为“缺席”",
    "信用管理：缺席一次扣除 10 分信用分，信用分低于 60 分则禁止继续预约",
    "统计查看：所有用户可查看时段利用率、热门座位排行和信用排行",
    "管理后台：管理员可增删自习室和座位，查看违约记录",
], 1):
    add_para(f"（{i}）{step}", indent=True)

add_para("技术特点：", bold=True)
add_table(
    ["特点", "说明"],
    [
        ["并发控制", "SELECT ... FOR UPDATE 行级锁防止多人同时预约同一座位"],
        ["事务管理", "预约、签到等关键操作使用数据库事务保证原子性"],
        ["定时任务", "MySQL 事件调度器每 5 分钟扫描超时预约并自动处理"],
        ["存储过程", "3 个核心存储过程封装预约、签到、取消业务逻辑"],
        ["数据库视图", "4 个视图提供不同维度的统计数据"],
        ["索引优化", "复合索引加速座位可用性查询和个人预约查询"],
        ["密码安全", "bcrypt 哈希加密，12 轮加盐"],
    ]
)

# ============================================================
# 二、需求分析
# ============================================================
page_break()
add_heading_styled("二、需求分析", 1)

add_heading_styled("1. 需求分析", 2)
add_para("1.1 功能需求", bold=True)
add_table(
    ["模块", "功能点", "详细描述"],
    [
        ["用户管理", "注册", "填写用户名、密码、真实姓名、学号，检查用户名和学号唯一性"],
        ["用户管理", "登录/登出", "bcrypt 密码验证，Flask session 管理登录状态"],
        ["用户管理", "角色区分", "student（学生）和 admin（管理员）两种角色，路由级权限控制"],
        ["座位浏览", "楼栋选择", "3 栋楼（图书馆主馆、教学楼A座、综合学习中心）"],
        ["座位浏览", "自习室浏览", "5 个自习室，显示楼层、容量等基本信息"],
        ["座位浏览", "座位地图", "每个自习室按 10 列网格展示全部座位"],
        ["座位浏览", "日期时段筛选", "日期选择器 + 7 个时段下拉框"],
        ["座位浏览", "状态标识", "绿色=可预约，红色=已占用，灰色=维护中，图标标注电源/靠窗"],
        ["预约管理", "预约座位", "检查信用分、座位状态、并发冲突后创建预约"],
        ["预约管理", "签到", "15 分钟内确认到达，自动更新状态"],
        ["预约管理", "取消预约", '仅可取消"已预约"状态的记录'],
        ["预约管理", "超时处理", "超过 15 分钟未签到自动标记缺席并扣分"],
        ["信用管理", "信用分", "默认 100 分，缺席扣 10 分"],
        ["信用管理", "预约限制", "信用分低于 60 分禁止预约"],
        ["信用管理", "违约记录", "每次缺席自动生成违约记录并关联扣分"],
        ["数据统计", "时段利用率", "近 30 天各时段预约数和使用率"],
        ["数据统计", "热门座位", "近 30 天被预约次数最多的座位 Top 20"],
        ["数据统计", "信用排行", "所有学生信用分排序，含缺席次数"],
        ["管理后台", "自习室管理", "添加新的自习室（楼栋、名称、楼层、容量）"],
        ["管理后台", "座位管理", "添加、删除座位，切换座位维护/可用状态"],
        ["管理后台", "违约查看", "查看最近 50 条违约记录"],
    ]
)

add_para("1.2 非功能需求", bold=True)
add_table(
    ["类型", "要求", "实现方式"],
    [
        ["并发安全", "同一座位同时段仅一人预约成功", "InnoDB 行级锁 + 事务"],
        ["数据一致性", "预约相关操作原子性", "手动事务管理（commit/rollback）"],
        ["查询效率", "常用查询响应 < 100ms", "3 个复合索引"],
        ["数据完整性", "实体、参照、用户定义完整性", "主键、外键、ENUM、UNIQUE 约束"],
        ["安全性", "密码不可逆存储", "bcrypt 哈希 12 轮加盐"],
        ["可维护性", "业务逻辑清晰分层", "MVC 架构（Flask 蓝图 + 独立模板）"],
    ]
)

add_heading_styled("2. 系统逻辑模型", 2)

add_para("2.1 数据流图（DFD）", bold=True)
add_para("顶层数据流图：")
add_para("用户 → 注册/登录 → 系统 ← 查看统计 ← 管理员")
add_para("  ↓                  ↓")
add_para("预约座位            管理座位")
add_para("  ↓                  ↓")
add_para("签到确认            处理违约")

add_para("一级细化 — 预约流程：")
add_para("学生 → [选择日期+时段] → [浏览可用座位] → [点击预约]")
add_para("  ↓ 存储过程检查信用分")
add_para("  ↓ SELECT ... FOR UPDATE 检查冲突")
add_para("  ├─ 有冲突 → ROLLBACK → “已被他人预约”")
add_para("  └─ 无冲突 → INSERT → COMMIT → “预约成功”")

add_para("2.2 数据字典", bold=True)
add_table(
    ["数据项", "类型", "取值范围", "说明"],
    [
        ["用户ID", "INT", "自增正整数", "唯一标识，主键"],
        ["用户名", "VARCHAR(50)", "字母数字组合", "唯一，用于登录"],
        ["密码", "VARCHAR(255)", "bcrypt 哈希", "不可逆加密"],
        ["信用分", "INT", "0-100，默认100", "低于60禁止预约"],
        ["角色", "ENUM", "student, admin", "控制权限"],
        ["座位标签", "VARCHAR(20)", "如 A-01", "自习室内部编号"],
        ["座位状态", "ENUM", "available, maintenance", "维护中不可预约"],
        ["预约状态", "ENUM", "5种状态", "reserved / checked_in / completed / cancelled / absent"],
        ["时间段", "TIME", "08:00-22:00", "每天7个时段，每段2小时"],
    ]
)

# ============================================================
# 三、系统设计
# ============================================================
page_break()
add_heading_styled("三、系统设计", 1)

add_heading_styled("1. 概念结构设计（E-R模型）", 2)

add_para("1.1 实体及其属性", bold=True)
add_para("用户（User）：{ id, username, password, real_name, student_id, credit, role, created_at }")
add_para("楼栋（Building）：{ id, name, location }")
add_para("自习室（Room）：{ id, name, floor, capacity, description }")
add_para("座位（Seat）：{ id, seat_label, has_power, near_window, status }")
add_para("时间段（TimeSlot）：{ id, slot_name, start_time, end_time }")
add_para("预约（Reservation）：{ id, reserve_date, status, created_at, checked_in_at }")
add_para("违约记录（Violation）：{ id, reason, points_deducted, created_at }")

add_para("1.2 实体间联系", bold=True)
add_table(
    ["联系名", "参与实体", "联系类型", "说明"],
    [
        ["位于", "楼栋 — 自习室", "1:N", "一个楼栋包含多个自习室"],
        ["容纳", "自习室 — 座位", "1:N", "一个自习室包含多个座位"],
        ["预约座位", "用户 — 座位", "M:N（通过预约）", "一个用户可预约多个座位"],
        ["按时段", "预约 — 时间段", "N:1", "每条预约属于一个时段"],
        ["产生", "预约 — 违约记录", "1:1", "一次缺席产生一条违约记录"],
    ]
)

add_para("1.3 E-R图", bold=True)
add_para("[楼栋] 1──N [自习室] 1──N [座位]")
add_para("                          │")
add_para("                          │ 1")
add_para("                          │")
add_para("[用户] 1──N [预约] N──1 [时间段]")
add_para("                │")
add_para("                │ 1")
add_para("                │")
add_para("            [违约记录]")

add_heading_styled("2. 逻辑结构设计", 2)

add_para("2.1 关系模式（7张表）", bold=True)

add_para("① user（用户表）", bold=True)
add_code("user(id INT PRIMARY KEY AUTO_INCREMENT,\n     username VARCHAR(50) NOT NULL UNIQUE,\n     password VARCHAR(255) NOT NULL,\n     real_name VARCHAR(50) NOT NULL,\n     student_id VARCHAR(20) UNIQUE,\n     credit INT DEFAULT 100,\n     role ENUM('student','admin') DEFAULT 'student',\n     created_at DATETIME DEFAULT CURRENT_TIMESTAMP)")
add_para("主码：id    候选码：username, student_id")

add_para("② building（楼栋表）", bold=True)
add_code("building(id INT PRIMARY KEY AUTO_INCREMENT,\n         name VARCHAR(100) NOT NULL,\n         location VARCHAR(200))")

add_para("③ room（自习室表）", bold=True)
add_code("room(id INT PRIMARY KEY AUTO_INCREMENT,\n     building_id INT NOT NULL,\n     name VARCHAR(100) NOT NULL,\n     floor INT,\n     capacity INT NOT NULL,\n     description TEXT,\n     FOREIGN KEY (building_id) REFERENCES building(id))")

add_para("④ seat（座位表）", bold=True)
add_code("seat(id INT PRIMARY KEY AUTO_INCREMENT,\n     room_id INT NOT NULL,\n     seat_label VARCHAR(20) NOT NULL,\n     has_power BOOLEAN DEFAULT FALSE,\n     near_window BOOLEAN DEFAULT FALSE,\n     status ENUM('available','maintenance') DEFAULT 'available',\n     FOREIGN KEY (room_id) REFERENCES room(id),\n     UNIQUE KEY uk_room_seat (room_id, seat_label))")

add_para("⑤ time_slot（时间段表）", bold=True)
add_code("time_slot(id INT PRIMARY KEY AUTO_INCREMENT,\n          slot_name VARCHAR(30) NOT NULL,\n          start_time TIME NOT NULL,\n          end_time TIME NOT NULL)")

add_para("⑥ reservation（预约表）—— 核心业务表", bold=True)
add_code("reservation(id INT PRIMARY KEY AUTO_INCREMENT,\n            user_id INT NOT NULL,\n            seat_id INT NOT NULL,\n            reserve_date DATE NOT NULL,\n            time_slot_id INT NOT NULL,\n            status ENUM('reserved','checked_in',\n                        'completed','cancelled','absent'),\n            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,\n            checked_in_at DATETIME,\n            FOREIGN KEY (user_id) REFERENCES user(id),\n            FOREIGN KEY (seat_id) REFERENCES seat(id),\n            FOREIGN KEY (time_slot_id) REFERENCES time_slot(id),\n            INDEX idx_seat_date_slot (seat_id, reserve_date, time_slot_id),\n            INDEX idx_user_date (user_id, reserve_date),\n            INDEX idx_status (status))")

add_para("⑦ violation（违约记录表）", bold=True)
add_code("violation(id INT PRIMARY KEY AUTO_INCREMENT,\n          user_id INT NOT NULL,\n          reservation_id INT,\n          reason VARCHAR(200),\n          points_deducted INT DEFAULT 10,\n          created_at DATETIME DEFAULT CURRENT_TIMESTAMP,\n          FOREIGN KEY (user_id) REFERENCES user(id),\n          FOREIGN KEY (reservation_id) REFERENCES reservation(id))")

add_para("2.2 范式分析", bold=True)
add_table(
    ["范式", "满足情况", "说明"],
    [
        ["1NF", "✓", "所有字段均为原子值，无重复组"],
        ["2NF", "✓", "非主属性完全函数依赖于主码，无部分依赖"],
        ["3NF", "✓", "不存在传递函数依赖"],
        ["BCNF", "✓", "所有决定因素都是候选码"],
    ]
)

add_para("2.3 索引设计", bold=True)
add_table(
    ["索引名", "表", "字段", "类型", "用途"],
    [
        ["PRIMARY", "所有表", "id", "聚簇索引", "主键快速定位"],
        ["idx_seat_date_slot", "reservation", "(seat_id, reserve_date, time_slot_id)", "复合索引", "座位可用性查询"],
        ["idx_user_date", "reservation", "(user_id, reserve_date)", "复合索引", "“我的预约”查询"],
        ["idx_status", "reservation", "(status)", "单列索引", "事件调度器扫描超时预约"],
    ]
)

add_heading_styled("3. 系统功能模块图", 2)
for line in [
    "自习室座位预约系统",
    "├── 用户模块",
    "│   ├── 注册",
    "│   ├── 登录",
    "│   ├── 登出",
    "│   └── Session验证",
    "├── 座位预约模块",
    "│   ├── 浏览座位",
    "│   ├── 预约座位",
    "│   ├── 签到",
    "│   └── 取消预约",
    "├── 信用模块",
    "│   ├── 信用分查询",
    "│   ├── 违约记录",
    "│   └── 超时自动扣分",
    "├── 管理后台模块",
    "│   ├── 自习室增删",
    "│   ├── 座位增删/状态切换",
    "│   └── 违约记录查看",
    "└── 数据统计模块",
    "    ├── 时段利用率",
    "    ├── 热门座位排行",
    "    └── 信用排行",
]:
    add_para(line)

add_heading_styled("4. 其他设计图形工具", 2)
add_para("4.1 系统架构图", bold=True)
for line in [
    "浏览器 (HTML/CSS/JS + Jinja2 Templates)",
    "    │ HTTP",
    "Flask Web 应用层",
    "    ├── auth.py (登录/注册/登出)",
    "    ├── seat.py (预约/签到/取消/统计)",
    "    ├── admin.py (管理后台)",
    "    └── models/db.py (数据访问层, PyMySQL)",
    "    │ TCP/3306",
    "MySQL 8.0 数据库层",
    "    ├── 表 ×7",
    "    ├── 视图 ×4",
    "    ├── 存储过程 ×3",
    "    ├── 事件调度器 ×2",
    "    └── 索引 ×3",
]:
    add_para(line)

add_para("4.2 预约状态转换图", bold=True)
for line in [
    "[reserved] 已预约，待签到",
    "    ├── 15min内签到 → [checked_in] 已签到 → 时段结束 → [completed] 已完成",
    "    ├── 取消预约 → [cancelled] 已取消",
    "    └── 超时未签到 → [absent] 缺席 → 扣除10分 → [violation] 违约记录",
]:
    add_para(line)

# ============================================================
# 四、详细设计
# ============================================================
page_break()
add_heading_styled("四、详细设计", 1)

add_heading_styled("4.1 数据库对象详细设计", 2)

add_para("4.1.1 存储过程", bold=True)
add_para("sp_create_reservation —— 预约座位：", bold=True)
add_code("""CREATE PROCEDURE sp_create_reservation(
    IN  p_user_id INT, IN  p_seat_id INT,
    IN  p_reserve_date DATE, IN  p_time_slot_id INT,
    OUT p_result INT, OUT p_message VARCHAR(100))
BEGIN
    START TRANSACTION;
    -- Step 1: 验证信用分 (>= 60)
    -- Step 2: 验证座位状态 (非 maintenance)
    -- Step 3: SELECT ... FOR UPDATE 加行锁检查冲突
    -- Step 4: 无冲突则 INSERT, COMMIT; 有冲突则 ROLLBACK
END""")

add_para("sp_check_in —— 签到：", bold=True)
add_code("""CREATE PROCEDURE sp_check_in(
    IN  p_reservation_id INT, IN  p_user_id INT,
    OUT p_result INT, OUT p_message VARCHAR(100))
BEGIN
    -- 1. 验证预约归属（是否为本人）
    -- 2. 验证状态（必须为 'reserved'）
    -- 3. 检查是否超时（15分钟窗口）
    -- 4. 未超时 -> UPDATE status = 'checked_in'
    -- 5. 已超时 -> 标记absent + 插入violation + credit-10
END""")

add_para("sp_cancel_reservation —— 取消预约：", bold=True)
add_code("""CREATE PROCEDURE sp_cancel_reservation(
    IN  p_reservation_id INT, IN  p_user_id INT,
    OUT p_result INT, OUT p_message VARCHAR(100))
BEGIN
    -- 1. 验证预约归属（是否为本人）
    -- 2. 验证状态（必须为 'reserved'）
    -- 3. UPDATE status = 'cancelled'
END""")

add_para("4.1.2 数据库事件（定时任务）", bold=True)
add_para("evt_auto_mark_absent —— 超时自动标记（每5分钟执行）：", bold=True)
add_para("扫描所有 status='reserved' 且距创建时间超过 15 分钟的预约记录，逐条标记为 'absent'，插入违约记录，扣除用户信用分 10 分（最低 0 分）。使用游标逐条处理。", indent=True)

add_para("evt_complete_reservations —— 自动完成（每天执行）：", bold=True)
add_para("将 status='checked_in' 且 reserve_date < 今天的预约记录批量更新为 'completed'。", indent=True)

add_para("4.1.3 数据库视图", bold=True)
add_table(
    ["视图名", "SQL 核心逻辑", "用途"],
    [
        ["v_seat_availability", "seat CROSS JOIN time_slot\nLEFT JOIN reservation\nON ... status IN ('reserved','checked_in')", "笛卡尔积+左连接，展示每个座位所有时段的占用情况"],
        ["v_slot_usage_stats", "time_slot LEFT JOIN reservation\nWHERE reserve_date >= 近30天\nGROUP BY slot_id", "按时间段分组统计利用率，按利用率降序"],
        ["v_seat_popularity", "seat → room → building\nLEFT JOIN reservation\nGROUP BY seat_id\nORDER BY COUNT(r.id) DESC", "三表连接+聚合，预约次数降序排列"],
        ["v_user_credit_ranking", "user LEFT JOIN reservation\nWHERE role='student'\nGROUP BY user_id\nORDER BY credit DESC", "学生信用排名，含总预约数和缺席数"],
    ]
)

add_para("4.1.4 应用层数据库访问设计", bold=True)
add_para("由于 PyMySQL 对存储过程的 OUT 参数支持不稳定，应用层实际采用“直接 SQL + Python 手动事务管理”方式：", indent=True)
add_code("""# 预约流程（controllers/seat.py）
conn = get_db()
try:
    with conn.cursor() as cur:
        # 1. 检查信用分
        cur.execute('SELECT credit FROM user WHERE id = %s', (user_id,))
        # 2. 检查座位状态
        cur.execute('SELECT status FROM seat WHERE id = %s', (seat_id,))
        # 3. 行级锁检查冲突
        cur.execute('''SELECT id FROM reservation
            WHERE seat_id = %s AND reserve_date = %s
            AND time_slot_id = %s
            AND status IN ('reserved', 'checked_in')
            FOR UPDATE''')
        if cur.fetchone():
            conn.rollback()
            flash('该座位已被他人预约', 'error')
        else:
            cur.execute('INSERT INTO reservation (...) VALUES (...)')
            conn.commit()
            flash('预约成功！', 'success')
except Exception as e:
    conn.rollback()""")

add_para("4.1.5 Session验证机制", bold=True)
add_code("""@app.before_request
def verify_session_user():
    if session.get('user_id'):
        user = query('SELECT id FROM user WHERE id = %s',
                     (session['user_id'],), one=True)
        if not user:
            session.clear()  # 数据库重建后自动清 session""")

add_para("4.1.6 权限控制设计", bold=True)
add_code("""# 登录检查装饰器
def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('auth.login'))
        return f(*args, **kwargs)
    return decorated

# 管理员检查装饰器
def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if session.get('role') != 'admin':
            flash('需要管理员权限', 'error')
            return redirect(url_for('seat.index'))
        return f(*args, **kwargs)
    return decorated""")

# ============================================================
# 五、系统实现与测试
# ============================================================
page_break()
add_heading_styled("五、系统实现与测试", 1)

add_heading_styled("1. 开发平台和工具选择", 2)
add_table(
    ["类别", "选择", "版本", "选择理由"],
    [
        ["数据库", "MySQL", "8.0", "课程指定数据库，InnoDB 引擎支持事务、行锁和外键约束"],
        ["后端框架", "Flask", "3.1", "Python 轻量级 Web 框架，内置 Jinja2 模板引擎"],
        ["数据库驱动", "PyMySQL", "1.2", "纯 Python 实现，无需编译 MySQL C 扩展"],
        ["密码加密", "bcrypt", "5.0", "安全哈希算法，自动加盐，抗彩虹表攻击"],
        ["前端技术", "HTML/CSS/JS", "原生", "不引入 React/Vue 等框架，降低复杂度"],
        ["字体", "Google Fonts", "—", "Playfair Display (标题) + DM Sans (正文)"],
        ["操作系统", "Windows 11", "—", "开发环境"],
    ]
)

add_heading_styled("2. 系统测试", 2)
add_para("2.1 功能测试用例", bold=True)
add_para("用户模块：", bold=True)
add_table(
    ["编号", "测试项", "操作步骤", "预期结果"],
    [
        ["T01", "新用户注册", "填写完整信息 → 点击注册", "提示“注册成功”，跳转登录页"],
        ["T02", "重复用户名", "使用已存在用户名注册", "提示“用户名已被注册”"],
        ["T03", "必填项校验", "不填用户名直接注册", "提示“请填写所有必填字段”"],
        ["T04", "正确登录", "输入正确的用户名密码", "跳转首页，显示用户名"],
        ["T05", "错误密码登录", "输入错误密码", "提示“用户名或密码错误”"],
        ["T06", "登出", "点击“退出”", "跳转登录页，session清空"],
        ["T07", "Session自动失效", "数据库重建后刷新", "自动跳转登录页"],
    ]
)

add_para("座位浏览与预约模块：", bold=True)
add_table(
    ["编号", "测试项", "操作步骤", "预期结果"],
    [
        ["T08", "浏览座位地图", "登录后进入首页", "显示各楼栋/自习室/座位，颜色区分状态"],
        ["T09", "切换日期时段", "选择不同日期和时段", "座位状态随选择刷新"],
        ["T10", "预约空闲座位", "点绿色座位“预约”", "提示“预约成功！”"],
        ["T11", "预约维护座位", "点灰色座位", "提示“该座位正在维护中”"],
        ["T12", "并发预约", "两用户同时同座预约", "仅一人成功，另一提示“已被他人预约”"],
        ["T13", "签到", "预约后就“我的预约”签到", "提示“签到成功”，状态变为已签到"],
        ["T14", "超时签到", "15分钟后签到", "提示“签到超时”，标记缺席并扣分"],
        ["T15", "取消预约", "在“我的预约”点取消", "提示“取消成功”，座位恢复可用"],
    ]
)

add_para("信用与统计模块：", bold=True)
add_table(
    ["编号", "测试项", "操作步骤", "预期结果"],
    [
        ["T16", "信用不足", "信用调至50后预约", "提示“信用分不足”"],
        ["T17", "缺席扣分", "超时未签到", "信用分-10，违约记录+1"],
        ["T18", "时段利用率", "点“统计”", "显示7个时段利用率和进度条"],
        ["T19", "热门座位", "统计页面", "Top 20座位排行"],
        ["T20", "信用排行", "统计页面", "学生信用排行（绿/黄/红）"],
    ]
)

add_para("管理后台模块：", bold=True)
add_table(
    ["编号", "测试项", "操作步骤", "预期结果"],
    [
        ["T21", "管理员访问", "admin点“管理后台”", "显示管理和违约记录"],
        ["T22", "学生被拒", "普通学生访问/admin/", "提示“需要管理员权限”"],
        ["T23", "添加座位", "填座位信息→添加", "座位列表新增一条"],
        ["T24", "切换状态", "点“切换状态”", "座位在可用/维护间切换"],
        ["T25", "添加自习室", "填信息→添加", "新增自习室出现"],
    ]
)

add_para("2.2 并发测试", bold=True)
add_para("测试场景：两个用户（admin 和 wxr666）同时对同一座位（seat_id=1）、同一日期（2026-06-09）、同一时段（slot_id=1）发起预约。", indent=True)
add_para("测试结果：", bold=True)
add_bullet("窗口 A（admin）：预约成功，提示“预约成功！请在15分钟内签到”")
add_bullet("窗口 B（wxr666）：预约失败，提示“该座位已被他人预约”")
add_para("技术原理：MySQL InnoDB 引擎的 SELECT ... FOR UPDATE 对目标行加排他锁。当第一个事务持有锁时，第二个事务必须等待。第一个事务 COMMIT 后，第二个事务的查询返回已有记录（cur.fetchone() 不为 None），触发 ROLLBACK。", indent=True)

add_para("2.3 数据完整性测试", bold=True)
add_table(
    ["测试项", "操作", "结果"],
    [
        ["外键约束 — 删除有预约的用户", "DELETE FROM user WHERE id=1", "✗ 被外键约束阻止"],
        ["外键约束 — 插入不存在的自习室ID", "INSERT INTO seat (room_id,...) VALUES (999,...)", "✗ 被外键约束阻止"],
        ["ENUM 约束 — 非法预约状态", "INSERT INTO reservation (status,...) VALUES ('invalid',...)", "✗ 被 ENUM 约束阻止"],
        ["UNIQUE 约束 — 重复座位号", "INSERT INTO seat VALUES (1, 'A-01')", "✗ 被唯一约束阻止"],
        ["UNIQUE 约束 — 重复用户名", "注册已存在的用户名", "✗ 提示“用户名已被注册”"],
    ]
)

# ============================================================
# 六、课程设计总结
# ============================================================
page_break()
add_heading_styled("六、课程设计总结", 1)

add_heading_styled("收获与体会", 2)

add_para("1. 数据库设计全流程实践", bold=True)
add_para("从需求分析 → 概念设计（E-R图）→ 逻辑设计（关系模式+范式分析）→ 物理设计（索引+存储过程+视图），完整走通了数据库设计的标准流程。特别是对范式理论的理解不再停留在概念层面——在设计 reservation 表时，能直观感受到如果违反 3NF（如在 reservation 中冗余 room_id），会导致座位调换自习室时的大量级联更新。", indent=True)

add_para("2. 并发控制的真刀实枪", bold=True)
add_para("“两人同时抢一个座位”不是教科书上的空洞场景，而是真实发生的功能需求。通过 SELECT ... FOR UPDATE 解决了并发冲突，对事务的隔离级别（REPEATABLE READ）和 InnoDB 锁机制（行锁 vs 表锁、共享锁 vs 排他锁）有了实操经验。", indent=True)

add_para("3. 存储过程 vs 应用层 SQL 的取舍", bold=True)
add_para("首先设计了完整的存储过程方案，但在 PyMySQL 驱动层遇到了 OUT 参数兼容性问题，最终改为应用层直接 SQL + 手动事务。这个过程让我理解了“技术选型需要结合驱动/框架的实际行为，而非盲目依赖数据库特性”。两种方案各有优劣：存储过程适合做数据密集型运算，应用层 SQL 更适合与 Web 框架的请求生命周期配合。", indent=True)

add_para("4. 信用机制的闭环设计", bold=True)
add_para("信用分不是孤立字段，而是一个闭环：预约 → 签到 → 缺席扣分 → 低分禁止预约。这种“行为 → 后果 → 约束”的设计模式在实际业务中非常常见，体现在数据库层面就是触发器/事件的联动。", indent=True)

add_heading_styled("项目亮点", 2)
add_bullet("并发控制：SELECT ... FOR UPDATE 行级锁确保同一座位同一时段仅一人预约成功")
add_bullet("自动化处理：MySQL 事件调度器每 5 分钟扫描超时预约，自动标记缺席、记录违约、扣除信用分")
add_bullet("信用体系：闭环的信用打分机制，违规成本与行为约束联动")
add_bullet("数据统计：4 个数据视图从不同维度（时段、座位、用户）提供分析数据")
add_bullet("完整性保证：7 个外键、2 个 UNIQUE 约束、3 个 ENUM 约束、3 个复合索引")

add_heading_styled("不足与改进方向", 2)
add_table(
    ["不足", "改进方向"],
    [
        ["页面刷新交互", "引入 AJAX/Fetch 实现无刷新预约和实时更新"],
        ["预约通知", "引入邮件/短信提醒签到"],
        ["权限粒度", "细化为超级管理员、馆长、值班员等多角色"],
        ["移动端适配", "座位地图在手机上排列过密，需专门适配"],
        ["未来日期预约", "当前可预约未来任意日期，可限制为近 7 天"],
        ["单元测试", "使用 pytest 为控制器和数据库操作编写自动化测试"],
    ]
)

add_heading_styled("总结", 2)
add_para("本项目以自习室座位预约这一真实校园需求为背景，从数据库设计出发，完整实现了包含用户管理、座位地图、预约签到、信用管理、后台管理和数据统计的 Web 系统。在技术实现上，重点展现了关系型数据库的高级特性——行级锁的并发控制、存储过程的业务封装、事件调度器的自动化处理和视图的多维统计。通过本次课程设计，对数据库系统概论的核心内容——关系模型、范式理论、索引优化、事务与并发、存储过程与触发器——有了从理论到实践的完整认知。", indent=True)

# ============================================================
# 保存文件
# ============================================================
output_path = "C:/Users/18773/Desktop/seat-reservation/report/课程设计报告.docx"
doc.save(output_path)
print(f"Word 文档已生成: {output_path}")
